"""
Générateur de rapports Word pour le projet CTI
Génère des rapports professionnels à partir des analyses de menaces
"""

import os
import json
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from io import BytesIO
import base64
from collections import Counter
from scripts.utils.database import DatabaseManager
from utils.logger import CTILogger

class WordReportGenerator:
    def __init__(self, config_path='config/database.json'):
        """Initialise le générateur de rapports Word"""
        self.logger = CTILogger("Word_Generator")
        
        # Initialiser la connexion à la base de données
        try:
            self.db = DatabaseManager(config_path)
            self.logger.info("Connexion à la base de données établie")
        except Exception as e:
            self.logger.error(f"Erreur de connexion à la base : {e}")
            self.db = None
        
        # Configuration des couleurs et styles
        self.colors = {
            'critical': RGBColor(220, 53, 69),    # Rouge
            'high': RGBColor(255, 193, 7),        # Orange
            'medium': RGBColor(255, 235, 59),     # Jaune
            'low': RGBColor(76, 175, 80),         # Vert
            'primary': RGBColor(0, 123, 255),     # Bleu
            'dark': RGBColor(52, 58, 64),         # Gris foncé
            'light': RGBColor(248, 249, 250)      # Gris clair
        }
        
        # Mapping des niveaux de sévérité
        self.severity_icons = {
            'critical': '🔴',
            'high': '🟠',
            'medium': '🟡',
            'low': '🟢',
            'unknown': '⚪'
        }

    def generate_comprehensive_report(self, analysis_data=None, days_back=7, 
                                    include_charts=True, include_detailed_tables=True):
        """Génère un rapport Word complet"""
        try:
            self.logger.info("Début de génération du rapport Word complet")
            
            # Récupérer les données d'analyse si non fournies
            if not analysis_data:
                from scripts.generators.threat_analyzer import ThreatAnalyzer
                analyzer = ThreatAnalyzer()
                analysis_data = analyzer.analyze_threat_landscape(days_back=days_back)
            
            if not analysis_data:
                self.logger.error("Aucune donnée d'analyse disponible")
                return None
            
            # Créer le document
            doc = Document()
            
            # Configuration des styles
            self._setup_document_styles(doc)
            
            # Générer les sections du rapport
            self._add_title_page(doc, analysis_data)
            self._add_executive_summary(doc, analysis_data)
            self._add_risk_assessment(doc, analysis_data)
            self._add_threat_landscape(doc, analysis_data)
            self._add_cve_analysis(doc, analysis_data, include_detailed_tables)
            self._add_ioc_analysis(doc, analysis_data, include_detailed_tables)
            self._add_alert_analysis(doc, analysis_data)
            self._add_temporal_analysis(doc, analysis_data)
            self._add_recommendations(doc, analysis_data)
            self._add_appendix(doc, analysis_data)
            
            # Optionnel : Ajouter des graphiques
            if include_charts:
                self._add_charts_section(doc, analysis_data)
            
            # Sauvegarder le rapport
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_dir = "output/word_reports"
            os.makedirs(output_dir, exist_ok=True)
            
            filename = f"CTI_Threat_Report_{timestamp}.docx"
            filepath = os.path.join(output_dir, filename)
            
            doc.save(filepath)
            
            self.logger.info(f"Rapport Word généré avec succès: {filepath}")
            return filepath
            
        except Exception as e:
            self.logger.error(f"Erreur lors de la génération du rapport Word : {e}")
            return None

    def _setup_document_styles(self, doc):
        """Configure les styles du document"""
        try:
            # Style pour les titres principaux
            title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Calibri'
            title_style.font.size = Inches(0.25)
            title_style.font.bold = True
            title_style.font.color.rgb = self.colors['primary']
            
            # Style pour les sous-titres
            subtitle_style = doc.styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
            subtitle_style.font.name = 'Calibri'
            subtitle_style.font.size = Inches(0.18)
            subtitle_style.font.bold = True
            subtitle_style.font.color.rgb = self.colors['dark']
            
            # Style pour le texte normal
            normal_style = doc.styles['Normal']
            normal_style.font.name = 'Calibri'
            normal_style.font.size = Inches(0.12)
            
        except Exception as e:
            self.logger.warning(f"Erreur lors de la configuration des styles : {e}")

    def _add_title_page(self, doc, analysis_data):
        """Ajoute la page de titre"""
        # Titre principal
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("RAPPORT D'ANALYSE CTI")
        run.font.size = Inches(0.35)
        run.font.bold = True
        run.font.color.rgb = self.colors['primary']
        
        doc.add_paragraph()  # Espace
        
        # Sous-titre
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("PAYSAGE DES MENACES CYBERNÉTIQUES")
        run.font.size = Inches(0.22)
        run.font.bold = True
        run.font.color.rgb = self.colors['dark']
        
        doc.add_paragraph()
        
        # Période d'analyse
        period = analysis_data.get('analysis_period', {})
        period_text = doc.add_paragraph()
        period_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = period_text.add_run(f"Période analysée: {period.get('days_analyzed', 'N/A')} jours")
        run.font.size = Inches(0.16)
        
        date_text = doc.add_paragraph()
        date_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_text.add_run(f"Du {period.get('start_date', 'N/A')[:10]} au {period.get('end_date', 'N/A')[:10]}")
        run.font.size = Inches(0.14)
        
        # Informations de génération
        doc.add_paragraph()
        doc.add_paragraph()
        
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ("Date de génération:", datetime.now().strftime('%d/%m/%Y %H:%M')),
            ("Éléments analysés:", str(analysis_data.get('data_summary', {}).get('total_collected_items', 0))),
            ("Niveau de risque:", analysis_data.get('risk_assessment', {}).get('risk_level', 'Unknown').upper()),
            ("Score de risque:", f"{analysis_data.get('risk_assessment', {}).get('overall_score', 0)}/100")
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = value
        
        # Saut de page
        doc.add_page_break()

    def _add_executive_summary(self, doc, analysis_data):
        """Ajoute le résumé exécutif"""
        # Titre de section
        heading = doc.add_heading("RÉSUMÉ EXÉCUTIF", level=1)
        heading.style.font.color.rgb = self.colors['primary']
        
        # Tableau de synthèse
        summary_table = doc.add_table(rows=1, cols=4)
        summary_table.style = 'Medium Grid 1 Accent 1'
        
        # En-têtes
        headers = ["Éléments collectés", "CVEs", "IoCs", "Alertes"]
        for i, header in enumerate(headers):
            cell = summary_table.cell(0, i)
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Données
        data_summary = analysis_data.get('data_summary', {})
        values = [
            str(data_summary.get('total_collected_items', 0)),
            str(data_summary.get('total_cves', 0)),
            str(data_summary.get('total_iocs', 0)),
            str(data_summary.get('total_alerts', 0))
        ]
        
        row = summary_table.add_row()
        for i, value in enumerate(values):
            row.cells[i].text = value
        
        doc.add_paragraph()
        
        # Texte de synthèse
        risk_assessment = analysis_data.get('risk_assessment', {})
        risk_level = risk_assessment.get('risk_level', 'unknown')
        risk_score = risk_assessment.get('overall_score', 0)
        
        # Paragraphe de risque avec mise en forme colorée
        risk_para = doc.add_paragraph()
        risk_para.add_run("NIVEAU DE RISQUE GLOBAL : ").font.bold = True
        
        risk_run = risk_para.add_run(f"{risk_level.upper()} ({risk_score}/100)")
        risk_run.font.bold = True
        
        # Couleur selon le niveau de risque
        if risk_level == 'critical':
            risk_run.font.color.rgb = self.colors['critical']
        elif risk_level == 'high':
            risk_run.font.color.rgb = self.colors['high']
        elif risk_level == 'medium':
            risk_run.font.color.rgb = self.colors['medium']
        else:
            risk_run.font.color.rgb = self.colors['low']
        
        doc.add_paragraph()
        
        # Points clés de l'analyse
        doc.add_paragraph("POINTS CLÉS :", style='Heading 2')
        
        # Menaces principales
        threat_categories = analysis_data.get('threat_categories', {})
        if threat_categories:
            sorted_threats = sorted(threat_categories.items(), key=lambda x: x[1]['count'], reverse=True)
            top_threats = sorted_threats[:3]
            
            threats_para = doc.add_paragraph()
            threats_para.add_run("• Principales catégories de menaces : ").font.bold = True
            threat_names = [f"{name} ({data['count']})" for name, data in top_threats if data['count'] > 0]
            threats_para.add_run(", ".join(threat_names))
        
        # CVEs critiques
        cve_analysis = analysis_data.get('cve_analysis', {})
        if cve_analysis.get('critical_cves'):
            cve_para = doc.add_paragraph()
            cve_para.add_run("• CVEs critiques identifiées : ").font.bold = True
            critical_count = len(cve_analysis['critical_cves'])
            cve_run = cve_para.add_run(f"{critical_count} vulnérabilités CVSS ≥ 9.0")
            if critical_count > 0:
                cve_run.font.color.rgb = self.colors['critical']
        
        # Alertes non résolues
        alert_analysis = analysis_data.get('alert_analysis', {})
        if alert_analysis.get('unresolved_count'):
            alert_para = doc.add_paragraph()
            alert_para.add_run("• Alertes non résolues : ").font.bold = True
            unresolved = alert_analysis['unresolved_count']
            alert_run = alert_para.add_run(f"{unresolved} alertes en attente")
            if unresolved > 5:
                alert_run.font.color.rgb = self.colors['high']
        
        doc.add_paragraph()

    def _add_risk_assessment(self, doc, analysis_data):
        """Ajoute l'évaluation des risques"""
        doc.add_heading("ÉVALUATION DES RISQUES", level=1)
        
        risk_assessment = analysis_data.get('risk_assessment', {})
        risk_factors = risk_assessment.get('risk_factors', {})
        
        # Tableau des facteurs de risque
        risk_table = doc.add_table(rows=1, cols=3)
        risk_table.style = 'Medium Grid 1 Accent 1'
        
        # En-têtes
        headers = ["Facteur de risque", "Score", "Impact"]
        for i, header in enumerate(headers):
            cell = risk_table.cell(0, i)
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Données des facteurs
        factor_names = {
            'critical_cves': 'CVEs critiques',
            'high_confidence_iocs': 'IoCs haute confiance',
            'unresolved_alerts': 'Alertes non résolues',
            'apt_activity': 'Activité APT',
            'ransomware_activity': 'Activité Ransomware'
        }
        
        for factor, score in risk_factors.items():
            if score > 0:
                row = risk_table.add_row()
                row.cells[0].text = factor_names.get(factor, factor)
                row.cells[1].text = str(score)
                
                # Impact selon le score
                if score >= 20:
                    impact = "🔴 Critique"
                elif score >= 10:
                    impact = "🟠 Élevé"
                elif score >= 5:
                    impact = "🟡 Modéré"
                else:
                    impact = "🟢 Faible"
                
                row.cells[2].text = impact
        
        # Recommandations prioritaires
        doc.add_paragraph()
        doc.add_paragraph("RECOMMANDATIONS PRIORITAIRES :", style='Heading 2')
        
        recommendations = risk_assessment.get('recommendations', [])
        for i, rec in enumerate(recommendations[:5], 1):
            rec_para = doc.add_paragraph()
            rec_para.add_run(f"{i}. ").font.bold = True
            rec_para.add_run(rec)

    def _add_threat_landscape(self, doc, analysis_data):
        """Ajoute l'analyse du paysage des menaces"""
        doc.add_heading("PAYSAGE DES MENACES", level=1)
        
        # Catégories de menaces
        doc.add_paragraph("Répartition par catégories :", style='Heading 2')
        
        threat_categories = analysis_data.get('threat_categories', {})
        if threat_categories:
            threat_table = doc.add_table(rows=1, cols=3)
            threat_table.style = 'Light List Accent 1'
            
            # En-têtes
            headers = ["Catégorie", "Occurrences", "Pourcentage"]
            for i, header in enumerate(headers):
                threat_table.cell(0, i).text = header
            
            total_threats = sum(data['count'] for data in threat_categories.values())
            
            for category, data in sorted(threat_categories.items(), 
                                       key=lambda x: x[1]['count'], reverse=True):
                if data['count'] > 0:
                    row = threat_table.add_row()
                    row.cells[0].text = category.replace('_', ' ').title()
                    row.cells[1].text = str(data['count'])
                    percentage = (data['count'] / total_threats * 100) if total_threats > 0 else 0
                    row.cells[2].text = f"{percentage:.1f}%"
        
        # Secteurs ciblés
        doc.add_paragraph()
        doc.add_paragraph("Secteurs ciblés :", style='Heading 2')
        
        targeted_sectors = analysis_data.get('targeted_sectors', {})
        if targeted_sectors:
            sectors_with_activity = {k: v for k, v in targeted_sectors.items() if v['count'] > 0}
            
            if sectors_with_activity:
                sector_table = doc.add_table(rows=1, cols=2)
                sector_table.style = 'Light List Accent 1'
                
                sector_table.cell(0, 0).text = "Secteur"
                sector_table.cell(0, 1).text = "Mentions"
                
                for sector, data in sorted(sectors_with_activity.items(), 
                                         key=lambda x: x[1]['count'], reverse=True):
                    row = sector_table.add_row()
                    row.cells[0].text = sector.replace('_', ' ').title()
                    row.cells[1].text = str(data['count'])
        
        # Mots-clés tendance
        doc.add_paragraph()
        doc.add_paragraph("Mots-clés tendance :", style='Heading 2')
        
        trending_keywords = analysis_data.get('trending_keywords', {})
        if trending_keywords:
            keywords_text = doc.add_paragraph()
            
            # Prendre les 10 premiers mots-clés
            top_keywords = list(trending_keywords.items())[:10]
            keyword_strings = [f"{word} ({count})" for word, count in top_keywords]
            
            keywords_text.add_run("• ")
            keywords_text.add_run(" • ".join(keyword_strings))

    def _add_cve_analysis(self, doc, analysis_data, include_details=True):
        """Ajoute l'analyse des CVEs"""
        doc.add_heading("ANALYSE DES VULNÉRABILITÉS (CVE)", level=1)
        
        cve_analysis = analysis_data.get('cve_analysis', {})
        
        if not cve_analysis or cve_analysis.get('total', 0) == 0:
            doc.add_paragraph("Aucune CVE identifiée dans la période analysée.")
            return
        
        # Statistiques générales
        doc.add_paragraph(f"Total des CVEs analysées : {cve_analysis.get('total', 0)}")
        
        # Distribution par sévérité
        severity_dist = cve_analysis.get('severity_distribution', {})
        if severity_dist:
            severity_table = doc.add_table(rows=1, cols=2)
            severity_table.style = 'Medium List 1 Accent 1'
            
            severity_table.cell(0, 0).text = "Niveau de sévérité"
            severity_table.cell(0, 1).text = "Nombre"
            
            severity_order = ['critical', 'high', 'medium', 'low', 'unknown']
            for severity in severity_order:
                count = severity_dist.get(severity, 0)
                if count > 0:
                    row = severity_table.add_row()
                    cell_0 = row.cells[0]
                    
                    # Ajouter l'icône et le texte
                    icon = self.severity_icons.get(severity, '')
                    cell_0.text = f"{icon} {severity.title()}"
                    row.cells[1].text = str(count)
        
        # CVEs critiques détaillées
        critical_cves = cve_analysis.get('critical_cves', [])
        if critical_cves and include_details:
            doc.add_paragraph()
            doc.add_paragraph("CVEs critiques (CVSS ≥ 9.0) :", style='Heading 2')
            
            for cve in critical_cves[:10]:  # Limiter à 10
                cve_para = doc.add_paragraph()
                cve_para.add_run(f"• {cve.get('cve_id', 'N/A')} ").font.bold = True
                cve_para.add_run(f"(CVSS: {cve.get('cvss_score', 'N/A')}) - ")
                
                description = cve.get('description', '')
                if len(description) > 200:
                    description = description[:200] + "..."
                cve_para.add_run(description)
        
        # Produits les plus affectés
        affected_products = cve_analysis.get('most_affected_products', {})
        if affected_products:
            doc.add_paragraph()
            doc.add_paragraph("Produits les plus affectés :", style='Heading 2')
            
            products_para = doc.add_paragraph()
            product_list = [f"{product} ({count})" for product, count in 
                          list(affected_products.items())[:5]]
            products_para.add_run("• " + " • ".join(product_list))

    def _add_ioc_analysis(self, doc, analysis_data, include_details=True):
        """Ajoute l'analyse des IoCs"""
        doc.add_heading("ANALYSE DES INDICATEURS DE COMPROMISSION", level=1)
        
        ioc_analysis = analysis_data.get('ioc_analysis', {})
        
        if not ioc_analysis or ioc_analysis.get('total', 0) == 0:
            doc.add_paragraph("Aucun IoC identifié dans la période analysée.")
            return
        
        # Statistiques générales
        doc.add_paragraph(f"Total des IoCs analysés : {ioc_analysis.get('total', 0)}")
        
        # Distribution par type
        type_dist = ioc_analysis.get('type_distribution', {})
        if type_dist:
            doc.add_paragraph("Répartition par type :", style='Heading 2')
            
            type_table = doc.add_table(rows=1, cols=2)
            type_table.style = 'Light List Accent 1'
            
            type_table.cell(0, 0).text = "Type d'IoC"
            type_table.cell(0, 1).text = "Nombre"
            
            for ioc_type, count in sorted(type_dist.items(), key=lambda x: x[1], reverse=True):
                row = type_table.add_row()
                row.cells[0].text = ioc_type
                row.cells[1].text = str(count)
        
        # Distribution par niveau de confiance
        confidence_dist = ioc_analysis.get('confidence_distribution', {})
        if confidence_dist:
            doc.add_paragraph()
            doc.add_paragraph("Répartition par niveau de confiance :", style='Heading 2')
            
            conf_para = doc.add_paragraph()
            conf_para.add_run(f"• Haute confiance (≥70%) : {confidence_dist.get('high', 0)}")
            conf_para = doc.add_paragraph()
            conf_para.add_run(f"• Confiance moyenne (30-70%) : {confidence_dist.get('medium', 0)}")
            conf_para = doc.add_paragraph()
            conf_para.add_run(f"• Faible confiance (<30%) : {confidence_dist.get('low', 0)}")
        
        # IoCs haute confiance
        high_confidence_iocs = ioc_analysis.get('high_confidence_iocs', [])
        if high_confidence_iocs and include_details:
            doc.add_paragraph()
            doc.add_paragraph("IoCs haute confiance (sélection) :", style='Heading 2')
            
            ioc_table = doc.add_table(rows=1, cols=4)
            ioc_table.style = 'Light Grid Accent 1'
            
            headers = ["Type", "Valeur", "Confiance", "Première détection"]
            for i, header in enumerate(headers):
                ioc_table.cell(0, i).text = header
            
            for ioc in high_confidence_iocs[:15]:  # Limiter à 15
                row = ioc_table.add_row()
                row.cells[0].text = ioc.get('type', 'N/A')
                
                # Tronquer la valeur si trop longue
                value = str(ioc.get('value', 'N/A'))
                if len(value) > 50:
                    value = value[:50] + "..."
                row.cells[1].text = value
                
                confidence = ioc.get('confidence', 0)
                row.cells[2].text = f"{confidence:.2f}"
                
                first_seen = ioc.get('first_seen', 'N/A')
                if isinstance(first_seen, str) and 'T' in first_seen:
                    first_seen = first_seen.split('T')[0]
                row.cells[3].text = str(first_seen)

    def _add_alert_analysis(self, doc, analysis_data):
        """Ajoute l'analyse des alertes"""
        doc.add_heading("ANALYSE DES ALERTES", level=1)
        
        alert_analysis = analysis_data.get('alert_analysis', {})
        
        if not alert_analysis or alert_analysis.get('total', 0) == 0:
            doc.add_paragraph("Aucune alerte identifiée dans la période analysée.")
            return
        
        # Statistiques générales
        total_alerts = alert_analysis.get('total', 0)
        unresolved = alert_analysis.get('unresolved_count', 0)
        resolution_rate = alert_analysis.get('resolution_rate', 0)
        
        stats_table = doc.add_table(rows=4, cols=2)
        stats_table.style = 'Medium List 1 Accent 1'
        
        stats_data = [
            ("Total des alertes", str(total_alerts)),
            ("Alertes non résolues", str(unresolved)),
            ("Alertes résolues", str(total_alerts - unresolved)),
            ("Taux de résolution", f"{resolution_rate:.1f}%")
        ]
        
        for i, (label, value) in enumerate(stats_data):
            stats_table.cell(i, 0).text = label
            stats_table.cell(i, 1).text = value
        
        # Distribution par sévérité
        severity_dist = alert_analysis.get('severity_distribution', {})
        if severity_dist:
            doc.add_paragraph()
            doc.add_paragraph("Répartition par sévérité :", style='Heading 2')
            
            sev_para = doc.add_paragraph()
            for severity, count in sorted(severity_dist.items(), key=lambda x: x[1], reverse=True):
                icon = self.severity_icons.get(severity, '')
                sev_para.add_run(f"• {icon} {severity.title()}: {count}  ")
        
        # Alertes critiques non résolues
        critical_unresolved = alert_analysis.get('critical_unresolved', [])
        if critical_unresolved:
            doc.add_paragraph()
            doc.add_paragraph("Alertes critiques non résolues :", style='Heading 2')
            
            for alert in critical_unresolved[:10]:
                alert_para = doc.add_paragraph()
                severity_icon = self.severity_icons.get(alert.get('severity', ''), '')
                alert_para.add_run(f"• {severity_icon} ").font.bold = True
                alert_para.add_run(f"{alert.get('title', 'Sans titre')} ")
                
                created_at = alert.get('created_at', 'N/A')
                if isinstance(created_at, str) and 'T' in created_at:
                    created_at = created_at.split('T')[0]
                alert_para.add_run(f"({created_at})")

    def _add_temporal_analysis(self, doc, analysis_data):
        """Ajoute l'analyse temporelle"""
        doc.add_heading("ANALYSE TEMPORELLE", level=1)
        
        temporal_analysis = analysis_data.get('temporal_analysis', {})
        
        if not temporal_analysis:
            doc.add_paragraph("Données d'analyse temporelle non disponibles.")
            return
        
        # Tendances des données collectées
        daily_stats = temporal_analysis.get('daily_collection_stats', {})
        if daily_stats:
            doc.add_paragraph("Évolution quotidienne des collectes :", style='Heading 2')
            
            # Tableau des statistiques quotidiennes
            temp_table = doc.add_table(rows=1, cols=4)
            temp_table.style = 'Light List Accent 1'
            
            headers = ["Date", "Items collectés", "CVEs", "IoCs"]
            for i, header in enumerate(headers):
                temp_table.cell(0, i).text = header
            
            # Trier par date et prendre les 7 derniers jours
            sorted_days = sorted(daily_stats.items())[-7:]
            for date, stats in sorted_days:
                row = temp_table.add_row()
                row.cells[0].text = date
                row.cells[1].text = str(stats.get('total_items', 0))
                row.cells[2].text = str(stats.get('cves', 0))
                row.cells[3].text = str(stats.get('iocs', 0))
        
        # Pics d'activité
        activity_peaks = temporal_analysis.get('activity_peaks', [])
        if activity_peaks:
            doc.add_paragraph()
            doc.add_paragraph("Pics d'activité détectés :", style='Heading 2')
            
            for peak in activity_peaks[:5]:
                peak_para = doc.add_paragraph()
                peak_para.add_run(f"• {peak.get('date', 'N/A')} : ").font.bold = True
                peak_para.add_run(f"{peak.get('count', 0)} éléments ")
                peak_para.add_run(f"({peak.get('type', 'N/A')})")
        
        # Analyse des tendances
        trends = temporal_analysis.get('trends', {})
        if trends:
            doc.add_paragraph()
            doc.add_paragraph("Tendances observées :", style='Heading 2')
            
            for trend_type, trend_data in trends.items():
                if trend_data.get('trend'):
                    trend_para = doc.add_paragraph()
                    trend_icon = "📈" if trend_data['trend'] == 'increasing' else "📉" if trend_data['trend'] == 'decreasing' else "➡️"
                    trend_para.add_run(f"• {trend_icon} {trend_type.replace('_', ' ').title()}: ")
                    trend_para.add_run(f"{trend_data['trend']} ({trend_data.get('change_percent', 0):.1f}%)")

    def _add_recommendations(self, doc, analysis_data):
        """Ajoute les recommandations"""
        doc.add_heading("RECOMMANDATIONS ET ACTIONS PRIORITAIRES", level=1)
        
        # Recommandations immédiates
        doc.add_paragraph("Actions immédiates requises :", style='Heading 2')
        
        immediate_actions = []
        
        # Basé sur les CVEs critiques
        cve_analysis = analysis_data.get('cve_analysis', {})
        critical_cves = cve_analysis.get('critical_cves', [])
        if len(critical_cves) > 0:
            immediate_actions.append(f"🔴 Évaluer et corriger {len(critical_cves)} CVEs critiques identifiées")
        
        # Basé sur les alertes non résolues
        alert_analysis = analysis_data.get('alert_analysis', {})
        unresolved_alerts = alert_analysis.get('unresolved_count', 0)
        if unresolved_alerts > 10:
            immediate_actions.append(f"🟠 Traiter {unresolved_alerts} alertes en attente de résolution")
        
        # Basé sur les IoCs haute confiance
        ioc_analysis = analysis_data.get('ioc_analysis', {})
        high_conf_iocs = ioc_analysis.get('high_confidence_iocs', [])
        if len(high_conf_iocs) > 50:
            immediate_actions.append(f"🟡 Analyser {len(high_conf_iocs)} IoCs haute confiance pour détection de compromission")
        
        # Afficher les actions immédiates
        for action in immediate_actions:
            doc.add_paragraph(action)
        
        if not immediate_actions:
            doc.add_paragraph("✅ Aucune action critique immédiate identifiée.")
        
        # Recommandations à moyen terme
        doc.add_paragraph()
        doc.add_paragraph("Recommandations à moyen terme :", style='Heading 2')
        
        medium_term_recs = [
            "📊 Améliorer la surveillance des secteurs les plus ciblés",
            "🔍 Renforcer la collecte d'intelligence sur les groupes APT actifs",
            "⚡ Optimiser les processus de réponse aux incidents",
            "📈 Développer des indicateurs de performance de sécurité (KPIs)",
            "🛡️ Mettre à jour les signatures et règles de détection"
        ]
        
        for rec in medium_term_recs:
            doc.add_paragraph(rec)
        
        # Recommandations stratégiques
        doc.add_paragraph()
        doc.add_paragraph("Orientations stratégiques :", style='Heading 2')
        
        strategic_recs = [
            "🎯 Développer des partenariats pour l'échange d'informations sur les menaces",
            "🔒 Renforcer la formation du personnel sur les nouvelles menaces",
            "🤖 Intégrer davantage d'automatisation dans les processus CTI",
            "📋 Établir des procédures de communication de crise",
            "🌐 Étendre la couverture géographique de la veille"
        ]
        
        for rec in strategic_recs:
            doc.add_paragraph(rec)

    def _add_appendix(self, doc, analysis_data):
        """Ajoute les annexes"""
        doc.add_heading("ANNEXES", level=1)
        
        # Méthodologie
        doc.add_paragraph("A. Méthodologie d'analyse", style='Heading 2')
        
        methodology_text = """
        Ce rapport est généré automatiquement à partir de l'analyse de multiples sources de renseignement sur les menaces cyber :
        
        • Sources de données : Flux RSS, APIs publiques, canaux Telegram, bases CVE/NVD
        • Période d'analyse : Données collectées sur les derniers jours spécifiés
        • Méthodes de scoring : Algorithmes propriétaires basés sur CVSS, confidence IoC, et corrélations
        • Mise à jour : Données actualisées en temps réel selon la fréquence des sources
        """
        
        doc.add_paragraph(methodology_text)
        
        # Sources de données
        doc.add_paragraph()
        doc.add_paragraph("B. Sources de données", style='Heading 2')
        
        data_summary = analysis_data.get('data_summary', {})
        sources_info = data_summary.get('sources_breakdown', {})
        
        if sources_info:
            sources_table = doc.add_table(rows=1, cols=2)
            sources_table.style = 'Light List Accent 1'
            
            sources_table.cell(0, 0).text = "Source"
            sources_table.cell(0, 1).text = "Éléments collectés"
            
            for source, count in sorted(sources_info.items(), key=lambda x: x[1], reverse=True):
                row = sources_table.add_row()
                row.cells[0].text = source
                row.cells[1].text = str(count)
        
        # Définitions et glossaire
        doc.add_paragraph()
        doc.add_paragraph("C. Glossaire", style='Heading 2')
        
        glossary = {
            "CVE": "Common Vulnerabilities and Exposures - Identifiant unique pour une vulnérabilité",
            "CVSS": "Common Vulnerability Scoring System - Système de notation des vulnérabilités (0-10)",
            "IoC": "Indicator of Compromise - Indicateur technique de compromission",
            "APT": "Advanced Persistent Threat - Menace persistante avancée",
            "CTI": "Cyber Threat Intelligence - Renseignement sur les menaces cyber",
            "TLP": "Traffic Light Protocol - Protocole de partage d'informations sensibles"
        }
        
        for term, definition in glossary.items():
            term_para = doc.add_paragraph()
            term_para.add_run(f"• {term} : ").font.bold = True
            term_para.add_run(definition)
        
        # Informations de contact
        doc.add_paragraph()
        doc.add_paragraph("D. Informations de contact", style='Heading 2')
        
        contact_text = """
        Pour toute question concernant ce rapport ou pour signaler des incidents :
        
        • Équipe CTI : cti-team@organization.com
        • SOC (24/7) : soc@organization.com
        • Urgences sécurité : +XX XXX XXX XXX
        
        Rapport généré automatiquement - Ne pas répondre directement à ce document.
        """
        
        doc.add_paragraph(contact_text)

    def _add_charts_section(self, doc, analysis_data):
        """Ajoute une section avec des graphiques (optionnel)"""
        try:
            doc.add_heading("GRAPHIQUES ET VISUALISATIONS", level=1)
            
            # Note : Cette méthode nécessiterait l'implémentation de génération de graphiques
            # avec matplotlib et leur insertion dans le document Word
            doc.add_paragraph("📊 Section réservée aux graphiques et visualisations.")
            doc.add_paragraph("(Implémentation requise pour la génération automatique de graphiques)")
            
        except Exception as e:
            self.logger.warning(f"Erreur lors de l'ajout des graphiques : {e}")

    def generate_executive_summary_only(self, analysis_data=None, days_back=7):
        """Génère uniquement un résumé exécutif court"""
        try:
            if not analysis_data:
                from scripts.generators.threat_analyzer import ThreatAnalyzer
                analyzer = ThreatAnalyzer()
                analysis_data = analyzer.analyze_threat_landscape(days_back=days_back)
            
            if not analysis_data:
                return None
            
            doc = Document()
            self._setup_document_styles(doc)
            
            # Titre simple
            title = doc.add_heading("RÉSUMÉ EXÉCUTIF CTI", level=1)
            title.style.font.color.rgb = self.colors['primary']
            
            # Période
            period = analysis_data.get('analysis_period', {})
            doc.add_paragraph(f"Période : {period.get('start_date', 'N/A')[:10]} au {period.get('end_date', 'N/A')[:10]}")
            
            # Résumé rapide
            self._add_executive_summary(doc, analysis_data)
            
            # Sauvegarder
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_dir = "output/word_reports"
            os.makedirs(output_dir, exist_ok=True)
            
            filename = f"CTI_Executive_Summary_{timestamp}.docx"
            filepath = os.path.join(output_dir, filename)
            
            doc.save(filepath)
            
            self.logger.info(f"Résumé exécutif généré : {filepath}")
            return filepath
            
        except Exception as e:
            self.logger.error(f"Erreur génération résumé exécutif : {e}")
            return None

    def get_analysis_stats(self):
        """Retourne des statistiques sur les analyses disponibles"""
        try:
            if not self.db:
                return {}
            
            # Requêtes pour obtenir des statistiques de base
            stats = {}
            
            # Statistiques sur les éléments collectés
            collected_query = """
                SELECT COUNT(*) as total, 
                       DATE(collected_at) as date
                FROM collected_items 
                WHERE collected_at >= NOW() - INTERVAL '7 days'
                GROUP BY DATE(collected_at)
                ORDER BY date DESC
            """
            
            # Statistiques CVE
            cve_query = """
                SELECT COUNT(*) as total,
                       severity,
                       DATE(collected_at) as date
                FROM cves 
                WHERE collected_at >= NOW() - INTERVAL '7 days'
                GROUP BY severity, DATE(collected_at)
                ORDER BY date DESC
            """
            
            # Statistiques IoCs
            ioc_query = """
                SELECT COUNT(*) as total,
                       ioc_type,
                       AVG(confidence_score) as avg_confidence
                FROM iocs 
                WHERE last_seen >= NOW() - INTERVAL '7 days'
                GROUP BY ioc_type
            """
            
            # Exécuter les requêtes (implémentation simplifiée)
            # Note: Nécessite l'implémentation complète avec self.db
            
            return stats
            
        except Exception as e:
            self.logger.error(f"Erreur récupération statistiques : {e}")
            return {}

# Fonction utilitaire pour test rapide
def main():
    """Fonction de test pour le générateur Word"""
    try:
        generator = WordReportGenerator()
        
        # Données de test simulées
        test_data = {
            'analysis_period': {
                'days_analyzed': 7,
                'start_date': '2024-07-14T00:00:00',
                'end_date': '2024-07-21T23:59:59'
            },
            'data_summary': {
                'total_collected_items': 1250,
                'total_cves': 45,
                'total_iocs': 320,
                'total_alerts': 28
            },
            'risk_assessment': {
                'risk_level': 'medium',
                'overall_score': 65,
                'risk_factors': {
                    'critical_cves': 15,
                    'high_confidence_iocs': 12,
                    'unresolved_alerts': 8
                }
            }
        }
        
        # Générer un rapport test
        result = generator.generate_comprehensive_report(analysis_data=test_data)
        
        if result:
            print(f"✅ Rapport de test généré : {result}")
        else:
            print("❌ Échec de génération du rapport test")
            
    except Exception as e:
        print(f"❌ Erreur lors du test : {e}")

if __name__ == "__main__":
    main()